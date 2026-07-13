"""
services/pdf_export.py — Converts a DOCX BytesIO buffer to PDF.

Uses python-docx to extract content and reportlab to render a clean PDF.
This is a server-side conversion so the frontend can offer PDF downloads
without any client-side dependencies.
"""
import io
import re
from datetime import date
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

# Parafinix colour palette
NAVY  = HexColor("#1F3346")
TEAL  = HexColor("#1F7A6C")
GREY  = HexColor("#5B6B79")
WHITE = HexColor("#FFFFFF")
RED   = HexColor("#C0392B")
LIGHT = HexColor("#EEF2F5")


def _styles():
    s = getSampleStyleSheet()
    return {
        "h1": ParagraphStyle("h1", fontSize=14, textColor=NAVY, spaceAfter=6,
                             spaceBefore=14, fontName="Helvetica-Bold"),
        "h2": ParagraphStyle("h2", fontSize=12, textColor=TEAL, spaceAfter=4,
                             spaceBefore=10, fontName="Helvetica-Bold"),
        "h3": ParagraphStyle("h3", fontSize=10.5, textColor=NAVY, spaceAfter=3,
                             spaceBefore=7, fontName="Helvetica-Bold"),
        "body": ParagraphStyle("body", fontSize=10, textColor=NAVY, spaceAfter=4,
                               spaceBefore=0, leading=14, alignment=TA_JUSTIFY),
        "small": ParagraphStyle("small", fontSize=8.5, textColor=GREY, spaceAfter=3,
                                leading=12),
        "centre": ParagraphStyle("centre", fontSize=10, textColor=NAVY, spaceAfter=4,
                                 alignment=TA_CENTER),
        "missing": ParagraphStyle("missing", fontSize=10, textColor=RED, spaceAfter=3,
                                  fontName="Helvetica-Bold"),
    }


def docx_to_pdf(docx_buf: io.BytesIO) -> io.BytesIO:
    """
    Converts a DOCX BytesIO to a PDF BytesIO.
    Extracts text and tables from the DOCX and re-renders in reportlab.
    """
    doc    = Document(docx_buf)
    styles = _styles()
    story  = []
    W      = A4[0] - 5 * cm

    # ── Parse paragraphs ──
    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            story.append(Spacer(1, 4))
            continue

        # Detect heading level from style name
        sn = (para.style.name or "").lower()
        if "heading 1" in sn or txt.isupper() and len(txt) < 80:
            story.append(Paragraph(txt, styles["h1"]))
        elif "heading 2" in sn:
            story.append(Paragraph(txt, styles["h2"]))
        elif "heading 3" in sn:
            story.append(Paragraph(txt, styles["h3"]))
        elif "[MISSING" in txt:
            story.append(Paragraph(txt, styles["missing"]))
        elif "PRIVATE AND CONFIDENTIAL" in txt or "Important Information" in txt:
            story.append(Paragraph(txt, styles["small"]))
        else:
            # Escape XML special chars for reportlab
            safe = txt.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, styles["body"]))

    # ── Parse tables ──
    for tbl in doc.tables:
        rows_data = []
        for i, row in enumerate(tbl.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_data)

        if not rows_data:
            continue

        num_cols = len(rows_data[0])
        col_w = [W / num_cols] * num_cols

        tbl_style = TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR",  (0, 0), (-1, 0), WHITE),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, -1), 9),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, LIGHT]),
            ("GRID",       (0, 0), (-1, -1), 0.5, GREY),
            ("VALIGN",     (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING",   (0, 0), (-1, -1), 6),
        ])

        # Wrap cell text in Paragraphs
        wrapped = []
        for ri, row in enumerate(rows_data):
            wr = []
            for cell in row:
                style = styles["small"] if ri == 0 else styles["body"]
                safe = cell.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                wr.append(Paragraph(safe, style))
            wrapped.append(wr)

        story.append(Spacer(1, 6))
        story.append(Table(wrapped, colWidths=col_w, style=tbl_style))
        story.append(Spacer(1, 8))

    # ── Render to PDF ──
    pdf_buf = io.BytesIO()
    pdf_doc = SimpleDocTemplate(
        pdf_buf, pagesize=A4,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
    )
    pdf_doc.build(story)
    pdf_buf.seek(0)
    return pdf_buf
