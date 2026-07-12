"""
documents/compliance_doc.py — Builds the separate Compliance Review
.docx: RAG rating, critical/flag/pass sections, full 28-point review
table, and the sign-off page.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
from datetime import date
from services.document_builder.word_builder import add_real_table, add_divider, add_heading, set_cell_bg

def build_compliance_doc(client_name, adviser_name, firm_name, report_ref, check_text, passes, flags, missing):
    doc = Document()
    for sec in doc.sections:
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    total = passes + flags + missing
    if total == 0: total = 1
    pct = int((passes / total) * 100)
    if missing > 5 or pct < 50:
        rag = 'RED'; rag_color = RGBColor(0xC0,0x39,0x2B); rag_bg = 'F8D7DA'
        rag_label = 'NOT READY FOR ISSUE — CRITICAL ITEMS OUTSTANDING'
    elif flags > 5 or missing > 0:
        rag = 'AMBER'; rag_color = RGBColor(0x85,0x64,0x04); rag_bg = 'FFF3CD'
        rag_label = 'REVIEW REQUIRED — ITEMS NEED ATTENTION BEFORE ISSUE'
    else:
        rag = 'GREEN'; rag_color = RGBColor(0x15,0x57,0x24); rag_bg = 'D4EDDA'
        rag_label = 'READY FOR ADVISER SIGN-OFF'

    # Cover
    for _ in range(2): doc.add_paragraph()
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(firm_name or 'YOUR FIRM NAME')
    fr.bold=True; fr.font.size=Pt(20); fr.font.color.rgb=RGBColor(0x1F,0x33,0x46)
    add_divider(doc)
    doc.add_paragraph()
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run('COMPLIANCE REVIEW REPORT')
    tr.bold=True; tr.font.size=Pt(18); tr.font.color.rgb=RGBColor(0x1F,0x33,0x46)
    doc.add_paragraph()
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run('FCA COBS 9A Suitability Assessment')
    sr.font.size=Pt(12); sr.font.color.rgb=RGBColor(0x5B,0x6B,0x79)
    for _ in range(2): doc.add_paragraph()

    add_real_table(doc,
        ['Field', 'Details'],
        [['Client', client_name or '[Not provided]'],
         ['Adviser', adviser_name or '[Not provided]'],
         ['Firm', firm_name or '[Not provided]'],
         ['Report reference', report_ref or f'SR-{date.today().strftime("%Y%m%d")}-001'],
         ['Review date', date.today().strftime('%d %B %Y')],
         ['Reviewed by', '[MISSING: Compliance reviewer name]'],
         ['Review standard', 'FCA COBS 9A — Suitability (30-point framework)']]
    )
    doc.add_paragraph()

    # RAG rating box
    rag_p = doc.add_paragraph(); rag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rag_r = rag_p.add_run(f'OVERALL COMPLIANCE RATING:  {rag}')
    rag_r.bold=True; rag_r.font.size=Pt(16); rag_r.font.color.rgb=rag_color
    label_p = doc.add_paragraph(); label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_r = label_p.add_run(rag_label)
    label_r.bold=True; label_r.font.size=Pt(11); label_r.font.color.rgb=rag_color
    doc.add_paragraph()

    # Score summary
    add_real_table(doc,
        ['Metric', 'Count', 'Percentage'],
        [['Items reviewed', str(total), '100%'],
         ['PASS', str(passes), f'{int(passes/total*100)}%'],
         ['FLAG — requires attention', str(flags), f'{int(flags/total*100)}%'],
         ['MISSING — not present', str(missing), f'{int(missing/total*100)}%'],
         ['Compliance score', f'{pct}%', rag]]
    )
    doc.add_page_break()

    # Parse compliance lines
    critical_items = []
    flag_items = []
    pass_items = []
    miss_items = []

    for line in check_text.split('\n'):
        if '|' not in line:
            continue
        parts = [p.strip() for p in line.split('|')]
        if len(parts) < 3:
            continue
        item_name = parts[0]
        status = parts[1] if len(parts) > 1 else ''
        finding = parts[2] if len(parts) > 2 else ''
        action = parts[3] if len(parts) > 3 else ''
        if 'MISSING' in status:
            critical_items.append([item_name, status, finding, action])
            miss_items.append([item_name, status, finding, action])
        elif 'FLAG' in status:
            flag_items.append([item_name, status, finding, action])
        elif 'PASS' in status:
            pass_items.append([item_name, status, finding, action])

    # Critical items section
    if miss_items or flag_items:
        add_heading(doc, 'SECTION 1: CRITICAL & FLAGGED ITEMS — ACTION REQUIRED', 1)
        add_divider(doc, 'C0392B')
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run('The following items must be addressed before this report can be issued to the client.')
        r.font.size=Pt(10.5); r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
        doc.add_paragraph()

        if miss_items:
            add_heading(doc, 'MISSING — Not Present in Report', 2)
            add_real_table(doc,
                ['#', 'Requirement', 'Finding', 'Required Action'],
                [[str(i+1), row[0], row[2], row[3]] for i, row in enumerate(miss_items)],
                'C0392B'
            )

        if flag_items:
            add_heading(doc, 'FLAG — Present but Incomplete or Unclear', 2)
            add_real_table(doc,
                ['#', 'Requirement', 'Finding', 'Required Action'],
                [[str(i+1), row[0], row[2], row[3]] for i, row in enumerate(flag_items)],
                'B8860B'
            )
        doc.add_page_break()

    # Full 30-point checklist
    add_heading(doc, 'SECTION 2: FULL COBS 9A COMPLIANCE CHECKLIST', 1)
    add_divider(doc)
    doc.add_paragraph()

    all_items = pass_items + flag_items + miss_items
    table_data = []
    counter = 1
    for item in check_text.split('\n'):
        if '|' not in item:
            continue
        parts = [p.strip() for p in item.split('|')]
        if len(parts) < 3:
            continue
        name = parts[0]; status = parts[1]
        finding = parts[2] if len(parts) > 2 else ''
        action = parts[3] if len(parts) > 3 else ''
        if name and status in ['PASS','FLAG','MISSING']:
            icon = '✓ PASS' if status=='PASS' else ('⚠ FLAG' if status=='FLAG' else '✗ MISSING')
            table_data.append([str(counter), name, icon, finding, action])
            counter += 1

    if table_data:
        add_real_table(doc,
            ['#', 'Requirement', 'Status', 'Finding', 'Required Action'],
            table_data
        )
    doc.add_page_break()

    # Paraplanner & adviser sign-off
    add_heading(doc, 'SECTION 3: SIGN-OFF AND AUTHORISATION', 1)
    add_divider(doc)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('This compliance review must be completed and signed off before the suitability report is issued to the client.').font.size = Pt(10.5)
    doc.add_paragraph()
    add_real_table(doc,
        ['Sign-off Stage', 'Name', 'Date', 'Signature / Initials', 'Notes'],
        [['Paraplanner review', '[Name]', '___/___/______', '_____________', ''],
         ['Compliance check', '[Name]', '___/___/______', '_____________', ''],
         ['Adviser approval', adviser_name or '[Name]', '___/___/______', '_____________', ''],
         ['Report issued to client', '[Name]', '___/___/______', '_____________', '']]
    )
    doc.add_paragraph()
    doc.add_paragraph()

    # Guidance notes
    add_heading(doc, 'SECTION 4: GUIDANCE NOTES', 2)
    for note in [
        'GREEN rating: All 30 items pass or only minor flags with clear resolutions. Report may proceed to adviser sign-off.',
        'AMBER rating: Some items flagged or missing. All flagged items must be resolved and re-reviewed before client issue.',
        'RED rating: Critical items missing. Report must not be issued until all MISSING items are resolved and a full re-review conducted.',
        'This compliance review is based on FCA COBS 9A suitability requirements and the FCA Consumer Duty (2023).',
        f'Review conducted: {date.today().strftime("%d %B %Y")} using Parafinix AI compliance framework.',
    ]:
        p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(note); r.font.size = Pt(10)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ── TEMPLATE EXTRACTION ───────────────────────────────────────

