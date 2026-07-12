"""
documents/factfind_doc.py — Builds a standalone Fact-Find document.

Three purposes in one document:
  A) Pre-filled form the adviser can give to the client to review/sign
  B) Internal working document showing what data exists and what is missing
  C) Gap analysis — clearly marked missing fields for follow-up

Structure:
  Page 1: Cover page (firm, adviser, client, date, status)
  Page 2: Instructions for client review
  Pages 3+: All 10 fact-find sections
    - Filled fields: shown clearly
    - Missing fields: shown as ____________ (blank line) in client version
                     shown as [MISSING] in red in internal version
  Final page: Client declaration and signature block
              Internal gap summary and action list
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import date
from services.document_builder.word_builder import (
    add_real_table, add_divider, add_heading, NAVY, TEAL, GREY, WHITE, RED_C, AMBER, GREEN
)


def _set_defaults(doc):
    for sec in doc.sections:
        sec.page_width    = Cm(21)
        sec.page_height   = Cm(29.7)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after  = Pt(4)
    style.paragraph_format.line_spacing = Pt(13.8)
    style.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT


def _body(doc, text, italic=False, color=None, size=Pt(10.5)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(13.8)
    r = p.add_run(text)
    r.font.size = size
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def _blank_line():
    """Returns a visual blank answer line for client-facing sections."""
    return '  _' + '_' * 55


def _val(value, missing_label='', client_facing=False):
    """
    Returns the field value for display.
    If missing: blank line for client-facing, [MISSING] for internal.
    """
    if value and value.strip() and value.strip() != '""':
        return value.strip()
    if client_facing:
        return _blank_line()
    return f'[MISSING{": " + missing_label if missing_label else ""}]'


def _section_heading(doc, title, icon=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f'{icon}  {title}' if icon else title)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '1F7A6C')
    pBdr.append(bot); pPr.append(pBdr)


def _field_row(doc, label, value, missing=False, client_facing=False):
    """Renders a label: value row as a two-column micro-table."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    row = table.rows[0]

    # Label cell
    lc = row.cells[0]
    lc.width = Cm(6)
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(2)
    lp.paragraph_format.space_after  = Pt(2)
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(10)
    lr.font.color.rgb = GREY

    # Value cell
    vc = row.cells[1]
    vc.width = Cm(10)
    vp = vc.paragraphs[0]
    vp.paragraph_format.space_before = Pt(2)
    vp.paragraph_format.space_after  = Pt(2)
    if missing and not client_facing:
        vr = vp.add_run(_val('', client_facing=False))
        vr.font.size = Pt(10)
        vr.font.color.rgb = RED_C
        vr.bold = True
    elif missing and client_facing:
        vr = vp.add_run(_blank_line())
        vr.font.size = Pt(10)
        vr.font.color.rgb = GREY
    else:
        vr = vp.add_run(str(value) if value else '')
        vr.font.size = Pt(10)


def build_factfind_doc(client_name: str, adviser_name: str, firm_name: str,
                       data: dict, client_facing: bool = False) -> io.BytesIO:
    """
    Builds the fact-find document.
    client_facing=True: blank lines for missing fields, suitable for client review
    client_facing=False: [MISSING] in red, suitable for internal/paraplanner use
    """
    doc = Document()
    _set_defaults(doc)

    p = data.get('personal', {})
    inc = data.get('income', {})
    exp = data.get('expenditure', {})
    assets = data.get('assets', {})
    pensions = data.get('pensions', [{}])
    protection = data.get('protection', [{}])
    liab = data.get('liabilities', {})
    obj = data.get('objectives', {})
    risk = data.get('risk', {})
    ep = data.get('estate_planning', {})
    tax = data.get('tax', {})
    flags = [f for f in data.get('flags', []) if f and f.strip()]

    doc_type = 'CLIENT REVIEW COPY' if client_facing else 'INTERNAL WORKING DOCUMENT'
    doc_color = '1F7A6C' if client_facing else 'C0392B'

    # ── COVER PAGE ───────────────────────────────────────────
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Cm(3)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(firm_name or 'YOUR FIRM NAME')
    cr.bold = True; cr.font.size = Pt(22); cr.font.color.rgb = NAVY

    _div_para(doc)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run('CLIENT FACT-FIND')
    tr.bold = True; tr.font.size = Pt(18); tr.font.color.rgb = NAVY
    tp.paragraph_format.space_before = Pt(10)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(doc_type)
    sr.bold = True; sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor.from_string(doc_color)
    sp.paragraph_format.space_before = Pt(6)

    pp2 = doc.add_paragraph()
    pp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp2r = pp2.add_run(f'Prepared for: {client_name or "[Client Name]"}')
    pp2r.bold = True; pp2r.font.size = Pt(13); pp2r.font.color.rgb = TEAL
    pp2.paragraph_format.space_before = Pt(14)

    doc.add_paragraph().paragraph_format.space_before = Pt(10)

    add_real_table(doc,
        ['Detail', 'Information'],
        [
            ['Adviser',       adviser_name or _val('', client_facing=client_facing)],
            ['Firm',          firm_name     or _val('', client_facing=client_facing)],
            ['Date prepared', date.today().strftime('%d %B %Y')],
            ['Status',        doc_type],
        ]
    )

    if client_facing:
        doc.add_paragraph().paragraph_format.space_before = Pt(16)
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        noter = note.add_run(
            'Please review the information below carefully.\n'
            'Correct any errors and complete any blank fields.\n'
            'Sign the declaration on the final page to confirm the information is accurate.'
        )
        noter.italic = True; noter.font.size = Pt(10); noter.font.color.rgb = GREY

    doc.add_page_break()

    # ── If internal: gap summary at top ──
    if not client_facing and flags:
        add_heading(doc, 'GAP ANALYSIS — ITEMS REQUIRING FOLLOW-UP', 1)
        gap_rows = [[str(i+1), flag] for i, flag in enumerate(flags)]
        add_real_table(doc, ['#', 'Missing or unclear item'], gap_rows, header_color='C0392B')
        doc.add_paragraph()
        _body(doc,
            f'{len(flags)} item(s) identified as missing or requiring clarification. '
            'These are highlighted in red throughout this document and must be '
            'obtained before a suitability report can be finalised.',
            color=RED_C)
        doc.add_page_break()

    # ── SECTION 1: PERSONAL DETAILS ──────────────────────────
    _section_heading(doc, '1. Personal Details', '👤')

    rows = []
    def r(label, key, hint=''):
        val = p.get(key, '')
        missing = not val or val.strip() == ''
        rows.append([label, _val(val, hint, client_facing) if missing else val])

    r('Client 1 full name',   'client1_name',   'Full legal name')
    r('Client 1 date of birth','client1_dob',   'DD/MM/YYYY')
    r('Client 2 full name',   'client2_name',   'If applicable')
    r('Client 2 date of birth','client2_dob',   'DD/MM/YYYY')
    r('Home address',         'address',        'Full address including postcode')
    r('Marital status',       'marital_status', 'Single/Married/Civil Partner/Divorced')
    r('Client 1 employment',  'client1_employment', 'Employed/Self-employed/Retired')
    r('Client 1 employer',    'client1_employer')
    r('Client 2 employment',  'client2_employment', 'If applicable')
    r('Client 2 employer',    'client2_employer')
    add_real_table(doc, ['Field', 'Information'], rows)

    deps = p.get('dependants', [])
    if deps and any(d.get('name') for d in deps):
        dep_rows = [[d.get('name',''), d.get('age',''), d.get('relationship','')] for d in deps if d.get('name')]
        add_real_table(doc, ['Dependant name', 'Age', 'Relationship'], dep_rows)
    else:
        _field_row(doc, 'Dependants',
                   _blank_line() if client_facing else '[MISSING: List all dependants]',
                   missing=True, client_facing=client_facing)

    # ── SECTION 2: INCOME ────────────────────────────────────
    _section_heading(doc, '2. Income', '💷')
    inc_rows = []
    def ir(label, key, hint=''):
        val = inc.get(key, '')
        missing = not val
        inc_rows.append([label, _val(val, hint, client_facing) if missing else val])
    ir('Client 1 gross salary (pa)', 'client1_gross_salary', 'Annual gross £')
    ir('Client 1 net monthly',       'client1_net_monthly',  'Take-home per month £')
    ir('Client 2 gross salary (pa)', 'client2_gross_salary', 'Annual gross £')
    ir('Client 2 net monthly',       'client2_net_monthly',  'Take-home per month £')
    ir('Client 1 bonus',             'client1_bonus',        '£ / % of salary')
    ir('Client 2 bonus',             'client2_bonus',        '£ / % of salary')
    ir('Other income sources',       'other_income',         'Rental, dividend, state pension etc.')
    ir('Salary sacrifice arrangement','salary_sacrifice',    'Pension/childcare/other')
    ir('Monthly surplus',            'monthly_surplus',      '£ per month after all expenditure')
    add_real_table(doc, ['Income item', 'Amount / Detail'], inc_rows)

    # ── SECTION 3: MONTHLY EXPENDITURE ───────────────────────
    _section_heading(doc, '3. Monthly Expenditure', '🧾')
    exp_rows = [
        ['Mortgage / rent', _val(exp.get('mortgage_payment',''), client_facing=client_facing) if not exp.get('mortgage_payment') else exp.get('mortgage_payment')],
        ['Total monthly expenditure', _val(exp.get('total_monthly',''), client_facing=client_facing) if not exp.get('total_monthly') else exp.get('total_monthly')],
        ['Notes', exp.get('notes','') or (_blank_line() if client_facing else '')],
    ]
    add_real_table(doc, ['Expenditure item', 'Amount'], exp_rows)

    # ── SECTION 4: ASSETS & SAVINGS ──────────────────────────
    _section_heading(doc, '4. Assets and Savings', '🏦')
    asset_rows = [
        ['Property value', _val(assets.get('property_value',''), 'Current market value', client_facing)],
        ['Outstanding mortgage', _val(assets.get('mortgage_balance',''), '£ balance', client_facing)],
        ['Property equity', _val(assets.get('property_equity',''), '£ estimated', client_facing)],
        ['Premium Bonds', assets.get('premium_bonds','') or (_blank_line() if client_facing else 'None noted')],
    ]
    add_real_table(doc, ['Asset', 'Value / Detail'], asset_rows)

    cash = assets.get('cash_savings', [])
    if cash and any(s.get('value') for s in cash):
        cash_rows = [[s.get('provider',''), s.get('type',''), s.get('value','')] for s in cash if s.get('value')]
        add_real_table(doc, ['Provider', 'Account type', 'Value'], cash_rows)
    else:
        _body(doc, 'Cash savings: ' + (_blank_line() if client_facing else '[MISSING: provider, type, value for each account]'))

    isas = assets.get('isas', [])
    if isas and any(i.get('value') for i in isas):
        isa_rows = [[i.get('provider',''), i.get('type',''), i.get('value',''), i.get('contributions','')] for i in isas if i.get('value')]
        add_real_table(doc, ['Provider', 'ISA type', 'Current value', 'Contributions'], isa_rows)
    else:
        _body(doc, 'ISA arrangements: ' + (_blank_line() if client_facing else '[MISSING: provider, type, value, contributions]'))

    gias = assets.get('gia', [])
    if gias and any(g.get('value') for g in gias):
        gia_rows = [[g.get('provider',''), g.get('value','')] for g in gias if g.get('value')]
        add_real_table(doc, ['GIA Provider', 'Current value'], gia_rows)

    # ── SECTION 5: PENSIONS ───────────────────────────────────
    _section_heading(doc, '5. Pension Arrangements', '🏛️')
    if pensions and any(pen.get('provider') or pen.get('current_value') for pen in pensions):
        pen_rows = []
        for pen in pensions:
            if pen.get('provider') or pen.get('current_value'):
                pen_rows.append([
                    pen.get('provider','') or _val('', 'Provider name', client_facing),
                    pen.get('type','')     or _val('', 'DC/DB/SIPP', client_facing),
                    pen.get('current_value','') or _val('', '£ current value', client_facing),
                    (pen.get('employee_contribution','') or '') + ' / ' + (pen.get('employer_contribution','') or ''),
                    pen.get('fund_name','') or _val('', 'Fund name', client_facing),
                    pen.get('selected_retirement_age','') or _val('', 'Target age', client_facing),
                    pen.get('status','') or 'Active',
                ])
        add_real_table(doc,
            ['Provider', 'Type', 'Current value', 'Contributions (EE/ER)', 'Fund', 'Target age', 'Status'],
            pen_rows)
    else:
        _body(doc, '[MISSING: All pension arrangement details required]', color=RED_C if not client_facing else None)

    # ── SECTION 6: PROTECTION ────────────────────────────────
    _section_heading(doc, '6. Protection Arrangements', '🛡️')
    if protection and any(pol.get('type') or pol.get('benefit') for pol in protection):
        prot_rows = []
        for pol in protection:
            if pol.get('type') or pol.get('benefit'):
                prot_rows.append([
                    pol.get('type','')    or _val('', 'Life/IP/CI', client_facing),
                    pol.get('provider','')or _val('', 'Provider', client_facing),
                    pol.get('benefit','') or _val('', '£ benefit', client_facing),
                    pol.get('premium','') or _val('', '£ pm', client_facing),
                    pol.get('term','')    or _val('', 'Years/to age', client_facing),
                    pol.get('basis','')   or _val('', 'Single/Joint', client_facing),
                ])
        add_real_table(doc,
            ['Policy type', 'Provider', 'Benefit', 'Premium pm', 'Term', 'Basis'],
            prot_rows)
    else:
        _body(doc, 'Protection arrangements: ' + (_blank_line() if client_facing else '[MISSING: list all policies]'))

    # ── SECTION 7: LIABILITIES ───────────────────────────────
    _section_heading(doc, '7. Mortgage and Liabilities', '🏠')
    liab_rows = [
        ['Mortgage balance',   liab.get('mortgage_balance','')  or _val('', '£ outstanding', client_facing)],
        ['Interest rate',      liab.get('mortgage_rate','')     or _val('', '% rate', client_facing)],
        ['Remaining term',     liab.get('mortgage_term','')     or _val('', 'Years remaining', client_facing)],
        ['Monthly payment',    liab.get('mortgage_monthly','')  or _val('', '£ per month', client_facing)],
        ['Other loans',        liab.get('other_loans','')       or (_blank_line() if client_facing else 'None noted')],
        ['Credit cards',       liab.get('credit_cards','')      or (_blank_line() if client_facing else 'None noted')],
    ]
    add_real_table(doc, ['Liability', 'Detail'], liab_rows)

    # ── SECTION 8: OBJECTIVES ────────────────────────────────
    _section_heading(doc, '8. Your Objectives', '🎯')
    obj_rows = [
        ['Primary objective',            obj.get('primary_objective','')             or _val('', 'State main objective', client_facing)],
        ['Target retirement age (C1)',   obj.get('target_retirement_age_client1','') or _val('', 'Age', client_facing)],
        ['Target retirement age (C2)',   obj.get('target_retirement_age_client2','') or _val('', 'Age if applicable', client_facing)],
        ['Target retirement income',     obj.get('target_retirement_income','')      or _val('', '£ pa net in today\'s terms', client_facing)],
        ['Secondary objectives',         obj.get('secondary_objectives','')          or _val('', 'List in priority order', client_facing)],
        ['Time horizons',                obj.get('time_horizons','')                 or _val('', 'Short/medium/long term', client_facing)],
        ['Key priorities',               obj.get('priorities','')                    or _val('', 'What matters most', client_facing)],
        ['Main concerns / worries',      obj.get('concerns','')                      or _val('', 'What keeps you up at night', client_facing)],
    ]
    add_real_table(doc, ['Objective / Question', 'Your Answer'], obj_rows)

    # ── SECTION 9: RISK ──────────────────────────────────────
    _section_heading(doc, '9. Attitude to Risk and Capacity for Loss', '📊')
    risk_rows = [
        ['Risk questionnaire used',       risk.get('questionnaire_tool','')         or _val('', 'Name of tool used', client_facing)],
        ['Risk questionnaire score',      risk.get('questionnaire_score','')        or _val('', 'Score from questionnaire', client_facing)],
        ['Risk profile assigned',         risk.get('profile_assigned','')           or _val('', 'e.g. Cautious/Balanced/Adventurous', client_facing)],
        ['Capacity for loss',             risk.get('capacity_for_loss','')          or _val('', 'Low/Medium/High + rationale', client_facing)],
        ['Investment time horizon',       risk.get('investment_time_horizon','')    or _val('', 'Years', client_facing)],
        ['Your own words on risk',        risk.get('client_own_words','')           or _val('', 'How you described your attitude to risk', client_facing)],
    ]
    add_real_table(doc, ['Risk assessment item', 'Detail'], risk_rows)

    # ── SECTION 10: ESTATE PLANNING ──────────────────────────
    _section_heading(doc, '10. Estate Planning', '📋')
    ep_rows = [
        ['Will in place?',          ep.get('will_in_place','')        or _val('', 'Yes/No', client_facing)],
        ['Will last reviewed',      ep.get('will_date','')            or _val('', 'Date of last review', client_facing)],
        ['LPA in place?',           ep.get('lpa_in_place','')         or _val('', 'Yes/No — Property & Finance / Health & Welfare', client_facing)],
        ['IHT position',            ep.get('iht_position','')         or _val('', 'Estimated estate value and NRB position', client_facing)],
        ['Beneficiaries noted',     ep.get('beneficiaries_noted','')  or _val('', 'Are beneficiaries up to date on all plans?', client_facing)],
        ['Trust arrangements',      ep.get('trust_arrangements','')   or (_blank_line() if client_facing else 'None noted')],
    ]
    add_real_table(doc, ['Estate planning item', 'Detail'], ep_rows)

    # ── SECTION 11: TAX POSITION ─────────────────────────────
    _section_heading(doc, '11. Tax Position', '💼')
    tax_rows = [
        ['Client 1 tax band',        tax.get('client1_tax_band','')      or _val('', 'Basic/Higher/Additional rate', client_facing)],
        ['Client 2 tax band',        tax.get('client2_tax_band','')      or _val('', 'Basic/Higher/Additional rate', client_facing)],
        ['CGT position',             tax.get('cgt_position','')          or _val('', 'Any carried-forward losses or gains', client_facing)],
        ['Annual allowance used',    tax.get('annual_allowance_used','') or _val('', '£ pension contributions this tax year', client_facing)],
        ['ISA allowance used',       tax.get('isa_allowance_used','')    or _val('', '£ ISA contributions this tax year', client_facing)],
        ['Tax notes',                tax.get('notes','')                 or (_blank_line() if client_facing else 'None noted')],
    ]
    add_real_table(doc, ['Tax item', 'Detail'], tax_rows)

    doc.add_page_break()

    # ── CLIENT DECLARATION (both versions) ───────────────────
    add_heading(doc, 'Declaration', 1)
    if client_facing:
        _body(doc,
            'I/We confirm that the information provided in this fact-find is accurate and complete '
            'to the best of my/our knowledge and belief. I/We understand that the financial advice '
            'provided will be based on this information, and I/We agree to notify my/our adviser '
            'promptly of any material changes to my/our circumstances.')
    else:
        _body(doc,
            'The client(s) should review, complete any blank sections, and sign below to confirm '
            'accuracy before this fact-find is used as the basis for a suitability report.')

    doc.add_paragraph()
    sig_rows = [
        ['Client 1 signature', '', 'Date', ''],
        ['Client 2 signature', '', 'Date', ''],
        ['Print name (C1)',    '', 'Print name (C2)', ''],
    ]
    sig_table = doc.add_table(rows=len(sig_rows), cols=4)
    sig_table.style = 'Table Grid'
    for ri, row_data in enumerate(sig_rows):
        for ci, cell_text in enumerate(row_data):
            cell = sig_table.rows[ri].cells[ci]
            p2 = cell.paragraphs[0]
            p2.paragraph_format.space_before = Pt(3)
            p2.paragraph_format.space_after  = Pt(16)
            r2 = p2.add_run(cell_text)
            r2.bold = (ci % 2 == 0)
            r2.font.size = Pt(10)

    # ── Internal version: gap action list ────────────────────
    if not client_facing:
        doc.add_paragraph()
        add_heading(doc, 'Paraplanner Action List — Items to Obtain', 1)
        if flags:
            for i, flag in enumerate(flags, 1):
                p3 = doc.add_paragraph()
                p3.paragraph_format.space_after = Pt(3)
                p3.paragraph_format.left_indent = Cm(0.5)
                r3 = p3.add_run(f'{i}. {flag}')
                r3.font.size = Pt(10.5)
        else:
            _body(doc, 'No gaps identified — all key fields are populated.', color=GREEN)

        doc.add_paragraph()
        _body(doc,
            f'Fact-find completed: {date.today().strftime("%d %B %Y")}  |  '
            f'Adviser: {adviser_name or "[MISSING]"}  |  '
            f'Firm: {firm_name or "[MISSING]"}  |  '
            'Generated by Parafinix AI — paraplanner review required before client issue.',
            color=GREY, size=Pt(9))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _div_para(doc):
    """Minimal divider paragraph for cover page."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '1F3346')
    pBdr.append(bot); pPr.append(pBdr)
