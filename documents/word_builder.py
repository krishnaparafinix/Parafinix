"""
documents/word_builder.py — Low-level Word rendering helpers.

Typography standards applied consistently to every document:
- Body text: Calibri 10.5pt, justified, 6pt after, 1.15 line spacing
- H1: Calibri Bold 14pt, navy, 16pt before, 4pt after, divider below
- H2: Calibri Bold 11.5pt, teal, 12pt before, 4pt after
- H3: Calibri Bold 10.5pt, navy italic, 8pt before, 2pt after
- List items: 10.5pt, 3pt after, hanging indent
- Table text: 10pt
- All paragraph text: JUSTIFIED (WD_ALIGN_PARAGRAPH.JUSTIFY)
- No orphan/widow paragraphs
"""
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── Colour palette ──────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x33, 0x46)
TEAL   = RGBColor(0x1F, 0x7A, 0x6C)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREY   = RGBColor(0x5B, 0x6B, 0x79)
RED_C  = RGBColor(0xC0, 0x39, 0x2B)
AMBER  = RGBColor(0xB8, 0x86, 0x0B)
GREEN  = RGBColor(0x2E, 0x7D, 0x52)

# ── Cell background ──────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ── Table builder ──────────────────────────────────────
def add_real_table(doc, header_row: list, data_rows: list, header_color: str = '1F3346'):
    """Builds a properly formatted Word table from lists of strings."""
    num_cols = len(header_row)
    if num_cols == 0:
        return

    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0]
    hdr.height = Pt(18)
    for i, cell_text in enumerate(header_row):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_color)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(str(cell_text).strip())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    for ri, row_data in enumerate(data_rows):
        row = table.rows[ri + 1]
        bg = 'FFFFFF' if ri % 2 == 0 else 'EEF2F5'
        for ci, cell_text in enumerate(row_data):
            if ci >= num_cols:
                break
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            txt = str(cell_text).strip()
            if '[MISSING' in txt:
                run = p.add_run(txt)
                run.font.size = Pt(10)
                run.font.color.rgb = RED_C
                run.bold = True
            else:
                _inline(p, txt, size=Pt(10))

# ── Markdown table parser ──────────────────────────────────────
def parse_md_tables(text: str) -> list:
    """
    Splits text into blocks: ('table', headers, rows) or ('text', line).
    Markdown tables (| col | col |) become real Word tables.
    """
    lines  = text.split('\n')
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('|') and line.strip().endswith('|') and len(line.strip()) > 2:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                header_row = [c.strip() for c in table_lines[0].strip('|').split('|')]
                data_start = 1
                if len(table_lines) > 1 and re.match(r'^[\|\s\-:]+$', table_lines[1].strip()):
                    data_start = 2
                data_rows = []
                for tl in table_lines[data_start:]:
                    row = [c.strip() for c in tl.strip('|').split('|')]
                    while len(row) < len(header_row):
                        row.append('')
                    data_rows.append(row[:len(header_row)])
                if any(h.strip() for h in header_row):
                    blocks.append(('table', header_row, data_rows))
            else:
                for tl in table_lines:
                    blocks.append(('text', tl))
        else:
            blocks.append(('text', line))
            i += 1
    return blocks

# ── Divider line ──────────────────────────────────────
def add_divider(doc, color: str = '1F3346', target_para=None):
    """
    Adds a bottom border. If target_para is given, applies the border
    directly to that paragraph (no new paragraph created).
    Otherwise creates a minimal 1pt spacer paragraph with a border.
    """
    if target_para is not None:
        pPr = target_para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), color)
        pBdr.append(bot)
        pPr.append(pBdr)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), color)
        pBdr.append(bot)
        pPr.append(pBdr)

# ── Heading builder ──────────────────────────────────────
def add_heading(doc, text: str, level: int = 1, page_break_before: bool = False):
    if page_break_before:
        doc.add_page_break()
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(8)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = NAVY
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "1F3346")
        pBdr.append(bot); pPr.append(pBdr)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = TEAL
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.bold = True
        r.italic = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = NAVY

# ── Inline text with bold / MISSING handling ──────────────────────────────────
def _inline(p, text: str, size = None):
    """Renders inline markdown bold (**text**) and [MISSING:] markers."""
    if size is None:
        size = Pt(11)
    parts = re.split(r'(\*\*[^*]+\*\*|\[MISSING:[^\]]*\])', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = size
        elif '[MISSING' in part:
            run = p.add_run(part)
            run.font.color.rgb = RED_C
            run.bold = True
            run.font.size = size
        else:
            run = p.add_run(part)
            run.font.size = size

# ── Body paragraph ──────────────────────────────────────
def _body_para(doc, text: str):
    """Standard justified body paragraph — the core prose style."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(14.3)
    _inline(p, text)
    return p

# ── Line renderer ──────────────────────────────────────
def render_line(doc, line: str, state: dict = None):
    """Converts a single line of AI markdown output into a Word element."""
    s = line.strip()

    if not s or s == '---':
        return

    if s.startswith('# '):
        pb = bool(state) and state.get('page_break_sections') and state.get('seen_h1')
        add_heading(doc, s[2:], 1, page_break_before=pb)
        if state is not None:
            state['seen_h1'] = True
        return

    if s.startswith('## '):
        add_heading(doc, s[3:], 2)
        return

    if s.startswith('### '):
        add_heading(doc, s[4:], 3)
        return

    if 'PARAPLANNER CHECK' in s.upper() and not s.startswith('|'):
        pb = bool(state) and state.get('page_break_sections') and state.get('seen_h1')
        add_heading(doc, 'PARAPLANNER CHECK — ITEMS FOR VERIFICATION BEFORE SIGN-OFF', 2, page_break_before=pb)
        if state is not None:
            state['seen_h1'] = True
        return

    if s.startswith('[MISSING'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(s)
        run.font.color.rgb = RED_C
        run.bold = True
        run.font.size = Pt(10.5)
        return

    if s.startswith('> '):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.line_spacing = Pt(13.8)
        run = p.add_run(s[2:])
        run.italic = True
        run.font.size = Pt(10.5)
        return

    if s.startswith('- ') or s.startswith('* '):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent    = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.line_spacing = Pt(13.8)
        p.add_run('•  ').bold = False
        _inline(p, s[2:])
        return

    if re.match(r'^\d+\.\s', s):
        content = re.sub(r'^\d+\.\s*', '', s)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent    = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.line_spacing = Pt(13.8)
        _inline(p, s)
        return

    _body_para(doc, s)


# ── Full content renderer ──────────────────────────────────────
def render_content(doc, text: str, page_break_sections: bool = False):
    """
    Renders an entire AI-generated markdown text block into the document.
    Tables become real Word tables.
    If page_break_sections=True, each new top-level section (# heading)
    and the Paraplanner Check start on a fresh page.
    """
    blocks = parse_md_tables(text)
    state = {'seen_h1': False, 'page_break_sections': page_break_sections}
    for block in blocks:
        if block[0] == 'table':
            _, header_row, data_rows = block
            add_real_table(doc, header_row, data_rows)
        else:
            _, line = block
            if not line.strip() or line.strip() == '---':
                continue
            render_line(doc, line, state)
