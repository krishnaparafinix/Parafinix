"""
documents/suitability_doc.py — Builds the full client-ready Suitability Report.

Typography is set ONCE in word_builder.py and applied consistently here.
The document structure:
  Page 1: Cover page
  Page 2+: Report content (via render_content)
  Final page: Important Information
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
from datetime import date
from docx.oxml.ns import qn as _qn
from docx.oxml import OxmlElement as _OxmlElement
from documents.word_builder import (
    add_real_table, add_divider, add_heading, render_content, NAVY, TEAL, GREY, WHITE
)



def _add_page_numbers(doc, client_name, firm_name):
    """Adds client name in header, page numbers in footer."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for sec in doc.sections:
        # Header: client name and report title, right-aligned
        sec.header.is_linked_to_previous = False
        if sec.header.paragraphs:
            hdr_p = sec.header.paragraphs[0]
        else:
            hdr_p = sec.header.add_paragraph()
        hdr_p.clear()
        hdr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hdr_r = hdr_p.add_run(f"{client_name or 'Client'} — Suitability Report")
        hdr_r.font.size = Pt(9)
        hdr_r.font.color.rgb = GREY
        hdr_r.italic = True

        # Footer: "Page X of Y  |  Firm name", centred
        if sec.footer.paragraphs:
            ftr_p = sec.footer.paragraphs[0]
        else:
            ftr_p = sec.footer.add_paragraph()
        ftr_p.clear()
        ftr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # "Page " plain text
        r1 = ftr_p.add_run("Page ")
        r1.font.size = Pt(9); r1.font.color.rgb = GREY

        # PAGE field
        fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
        r1._r.append(fld1)
        r2 = ftr_p.add_run()
        instr1 = OxmlElement('w:instrText'); instr1.text = 'PAGE'
        r2._r.append(instr1)
        fld1e = OxmlElement('w:fldChar'); fld1e.set(qn('w:fldCharType'), 'end')
        r2._r.append(fld1e)

        # " of " plain text
        r3 = ftr_p.add_run(" of ")
        r3.font.size = Pt(9); r3.font.color.rgb = GREY

        # NUMPAGES field
        fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'begin')
        r3._r.append(fld2)
        r4 = ftr_p.add_run()
        instr2 = OxmlElement('w:instrText'); instr2.text = 'NUMPAGES'
        r4._r.append(instr2)
        fld2e = OxmlElement('w:fldChar'); fld2e.set(qn('w:fldCharType'), 'end')
        r4._r.append(fld2e)

        # " | Firm name"
        r5 = ftr_p.add_run(f"  |  {firm_name or 'Parafinix AI'}")
        r5.font.size = Pt(9); r5.font.color.rgb = GREY



def _set_doc_defaults(doc):
    """Apply global document defaults once."""
    for sec in doc.sections:
        sec.page_width      = Cm(21)
        sec.page_height     = Cm(29.7)
        sec.top_margin      = Cm(2.5)
        sec.bottom_margin   = Cm(2.5)
        sec.left_margin     = Cm(2.5)
        sec.right_margin    = Cm(2.5)
    style = doc.styles['Normal']
    style.font.name        = 'Calibri'
    style.font.size        = Pt(10.5)
    style.paragraph_format.space_after  = Pt(6)
    style.paragraph_format.line_spacing = Pt(13.8)
    style.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY


def _cover_page(doc, client_name, adviser_name, firm_name, basis, charges, report_ref):
    """Adds a clean, professional cover page."""
    firm_p = doc.add_paragraph()
    firm_p.paragraph_format.space_before = Cm(3)
    firm_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firm_r = firm_p.add_run(firm_name or 'YOUR FIRM NAME')
    firm_r.bold = True
    firm_r.font.size = Pt(24)
    firm_r.font.color.rgb = NAVY
    firm_p.paragraph_format.space_after = Pt(6)

    add_divider(doc)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(14)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_r = title_p.add_run('SUITABILITY REPORT')
    title_r.bold = True
    title_r.font.size = Pt(20)
    title_r.font.color.rgb = NAVY
    title_p.paragraph_format.space_after = Pt(10)

    client_p = doc.add_paragraph()
    client_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    client_r = client_p.add_run(f'Prepared for: {client_name or "[Client Name]"}')
    client_r.bold = True
    client_r.font.size = Pt(14)
    client_r.font.color.rgb = TEAL
    client_p.paragraph_format.space_after = Pt(20)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    add_real_table(doc,
        ['Detail', 'Information'],
        [
            ['Adviser',          adviser_name or '[MISSING: Adviser name]'],
            ['Firm',             firm_name     or '[MISSING: Firm name]'],
            ['Basis of advice',  basis         or '[MISSING: Confirm basis]'],
            ['Date of report',   date.today().strftime('%d %B %Y')],
            ['Report reference', report_ref    or f'SR-{date.today().strftime("%Y%m%d")}-001'],
            ['FCA reference',    '[MISSING: FCA FRN — insert before issue]'],
            ['Status',           'SUBJECT TO PARAPLANNER AND ADVISER SIGN-OFF'],
        ]
    )

    conf_p = doc.add_paragraph()
    conf_p.paragraph_format.space_before = Pt(20)
    conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_r = conf_p.add_run(
        'PRIVATE AND CONFIDENTIAL\n'
        'This document is intended solely for the named recipient.\n'
        'Please read carefully and contact your adviser if anything requires clarification.'
    )
    conf_r.font.size = Pt(9)
    conf_r.italic = True
    conf_r.font.color.rgb = GREY

    doc.add_page_break()


def _important_information(doc, firm_name):
    """Adds the Important Information page."""
    add_heading(doc, 'IMPORTANT INFORMATION', 1)
    add_divider(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    items = [
        f'This report is based on the information you provided during your financial planning '
        f'review and is correct as at {date.today().strftime("%d %B %Y")}. If your circumstances '
        f'have changed since this date, please contact your adviser before acting on any '
        f'recommendation contained herein.',
        'The value of investments can fall as well as rise. You may get back less than you invest. '
        'Past performance is not a reliable guide to future performance.',
        'Tax treatment depends on individual circumstances and is subject to change. This report '
        'does not constitute legal or tax advice.',
        'Cash-flow projections and illustrative figures are based on stated assumptions and are '
        'not a guarantee of future outcomes.',
        'This report was prepared with AI assistance (Parafinix AI) and has been reviewed and '
        'approved by a qualified adviser before issue to the client.',
        f'{firm_name or "[Firm name]"} is authorised and regulated by the Financial Conduct '
        f'Authority. [MISSING: FCA FRN]',
    ]

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.line_spacing = Pt(13.8)
        run = p.add_run(item)
        run.font.size = Pt(9.5)
        run.font.color.rgb = GREY


def build_suitability_doc(client_name, adviser_name, firm_name, basis,
                           charges, report_ref, part1, part2, part3="", part4="") -> io.BytesIO:
    doc = Document()
    _set_doc_defaults(doc)
    _cover_page(doc, client_name, adviser_name, firm_name, basis, charges, report_ref)
    # Combine all three passes into one document
    full_content = part1
    if part2: full_content += "\n\n" + part2
    if part3: full_content += "\n\n" + part3
    if part4: full_content += "\n\n" + part4
    render_content(doc, full_content, page_break_sections=True)
    doc.add_page_break()
    _important_information(doc, firm_name)
    _add_page_numbers(doc, client_name, firm_name)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
